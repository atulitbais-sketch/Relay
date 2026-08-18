from __future__ import annotations

from io import BytesIO
from pathlib import Path


def extract_text(filename: str, raw: bytes) -> str:
    """Extract searchable text from supported document formats."""
    extension = Path(filename).suffix.lower()

    if extension in {".txt", ".md", ".csv", ".json", ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".css"}:
        return raw.decode("utf-8", errors="replace")

    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF support requires the pypdf package.") from exc
        reader = PdfReader(BytesIO(raw))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(pages).strip()

    if extension == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("DOCX support requires the python-docx package.") from exc
        document = DocxDocument(BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs).strip()

    raise ValueError(
        f"Unsupported file type '{extension or 'unknown'}'. Supported types: TXT, MD, PDF and DOCX."
    )


def format_size(byte_count: int) -> str:
    if byte_count < 1024:
        return f"{byte_count} B"
    if byte_count < 1024 * 1024:
        return f"{byte_count / 1024:.1f} KB"
    return f"{byte_count / (1024 * 1024):.1f} MB"
